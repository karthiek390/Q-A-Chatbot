import os
import re
import shutil
import pdfplumber
import numpy as np
import pandas as pd
from pathlib import Path
from docx import Document

class Preprocessing:
    
    def __init__(self, model, filepath):
        self.model = model
        self.file = filepath
        self.para_dict = {}
        
    def data_ingestion(self):
        file_path = os.path.dirname(self.file)
        file_name = os.path.basename(self.file)
        file, file_extension = os.path.splitext(file_name)  # Fix here
        file_extension = file_extension.lower()[1:]  # Remove the dot and make lowercase

        try:
            shutil.rmtree(os.path.join(file_path, "files"))
        except Exception:
            pass
        try:
            os.makedirs(os.path.join(file_path, "files"))
        except Exception:
            pass

        if file_extension == 'txt':
            document = Document()
            myfile = open(self.file).read()
            myfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', myfile)  # Remove all non-XML-compatible characters
            document.add_paragraph(myfile)
            document.save(os.path.join(file_path, "files", file + '.docx'))

        elif file_extension == 'docx':
            document = Document(self.file)
            document.save(os.path.join(file_path, "files", file_name))
        
        elif file_extension == 'pdf':
            with pdfplumber.open(self.file) as pdf:
                pages = pdf.pages
                document = Document()
                for i, page in enumerate(pages):
                    pdf_text = page.extract_text()
                    document.add_paragraph(pdf_text)
                document.save(os.path.join(file_path, "files", file + '.docx'))
            
        # Locate the newly created .docx file
        directory = os.path.join(file_path, "files")
        pathlist = Path(directory).rglob('*.docx')
        doc_file = ''
        for path in pathlist:
            doc_file = str(path)
        return doc_file
    
    def preprocessing(self):
        doc_file = self.data_ingestion()
        document = Document(doc_file)
        doc_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        text = doc_text.replace('\n', ' ').replace("~", "").replace("–", "")
        text = re.sub(r"[\[].*?[\]]", "", text)
        text = text.replace("“", "'").replace("”", "'").replace("’", "'")
        
        sentence_list = text.split(". ")
        sentence_list = list(filter(None, sentence_list))
        sentence_list = [i + "." if len(i) > 0 and i[-1] != "." else i for i in sentence_list]
        return sentence_list
    
    def doc_paragraphs(self):
        preprocessed_sentences = self.preprocessing()
        para_dict = {}
        count = 0
        index_num = 0
        while preprocessed_sentences:
            string = " "
            for i in range(len(preprocessed_sentences)):
                if len(string.split(" ")) + len(preprocessed_sentences[i].split(" ")) < 250:
                    string += preprocessed_sentences[i]
                    index_num = i
                else:
                    break
            preprocessed_sentences = preprocessed_sentences[index_num + 1:]  
            count += 1
            para_dict[count] = string
        return para_dict
        
    def auto_qna(self):
        para_dict = self.doc_paragraphs()
        qna_list = []
        for i in range(len(para_dict.keys())):
            try:
                qna_list.append(self.model(para_dict[i + 1]))
            except ValueError:
                continue

        qna_list = [item for sublist in qna_list for item in sublist]
        qna_df = pd.DataFrame(qna_list)
        return qna_df
    
    def main(self):
        df = self.auto_qna()
        return df
